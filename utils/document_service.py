from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import logging

logger = logging.getLogger(__name__)


def apply_run_style(run, font_size=12, bold=None):
    """
    Surgically applies typography styling to a run while preserving color and other styles.
    """
    run.font.name = 'Bookman Old Style'
    run.font.size = Pt(font_size)
    if bold is not None:
        run.bold = bold
    
    # Global font mapping enforcement via XML to ensure consistency across all renderers
    r = run._element.get_or_add_rPr()
    rFonts = r.get_or_add_rFonts()
    for attr in ['w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs']:
        rFonts.set(qn(attr), 'Bookman Old Style')


def safe_update_runs(paragraph, new_text):
    """
    Updates paragraph text by modifying existing runs to preserve formatting.
    """
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return

    # Update first run and clear others to maintain style of the first run
    paragraph.runs[0].text = new_text
    for i in range(1, len(paragraph.runs)):
        paragraph.runs[i].text = ""


def apply_academic_formatting(paragraph, is_heading=False):
    """
    Applies standard academic styling to non-table paragraphs without destroying existing styles.
    """
    # Global Spacing Rule: 1.5 for ALL paragraphs
    paragraph.paragraph_format.line_spacing = 1.5
    
    # Alignment Protection Logic:
    # If alignment is missing (None) or broken (Right), apply specific defaults.
    # Preserve if already defined (e.g., Center, Justify, Left).
    if paragraph.alignment in [None, WD_ALIGN_PARAGRAPH.RIGHT]:
        if is_heading:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    font_size = 14 if is_heading else 12
    is_bold = True if is_heading else False
    
    if is_heading:
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)

    # Apply style to all existing runs
    for run in paragraph.runs:
        apply_run_style(run, font_size=font_size, bold=is_bold)


def process_label_value_paragraph(paragraph):
    """
    Surgically formats Label : Value patterns without paragraph.clear().
    """
    text = paragraph.text
    if ":" not in text:
        for run in paragraph.runs:
            apply_run_style(run, font_size=12, bold=False)
        return

    parts = text.split(":", 1)
    label_part = parts[0].strip() + " :"
    value_part = " " + parts[1].strip()

    if not paragraph.runs:
        paragraph.add_run("")

    # Update first run for label
    first_run = paragraph.runs[0]
    original_color = first_run.font.color.rgb
    first_run.text = label_part
    apply_run_style(first_run, font_size=12, bold=True)
    if original_color:
        first_run.font.color.rgb = original_color

    # Update or add second run for value
    if len(paragraph.runs) > 1:
        second_run = paragraph.runs[1]
        second_run.text = value_part
        apply_run_style(second_run, font_size=12, bold=False)
        # Clear any extra runs
        for i in range(2, len(paragraph.runs)):
            paragraph.runs[i].text = ""
    else:
        v_run = paragraph.add_run(value_part)
        apply_run_style(v_run, font_size=12, bold=False)


def fill_reflective_journal(template_path: str, output_path: str, data: dict) -> str:
    doc = Document(template_path)

    # Standardize Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Data mapping according to Input Contract
    # Hard Locked Fields (from user input)
    student_details = data.get("student_details", {})
    gen_content = data.get("generated_content", {})
    
    mappings = {
        "name of the student": student_details.get("student_name", ""),
        "course": student_details.get("course_name", ""),
        "instructor": student_details.get("instructor", ""),
        "assessment": student_details.get("assessment", ""),
        "date": data.get("date", ""),
        "journal entry topic": data.get("journal_topic", ""),
        "experience": gen_content.get("experience", ""),
        "feelings": gen_content.get("feelings", ""),
        "learning": gen_content.get("learning", ""),
        "application": gen_content.get("application", ""),
        "conclusion": gen_content.get("conclusion", ""),
    }

    # 1. Fill data (Maintain mapping logic with strict leakage protection)
    priority_keys = [
        "journal entry topic", "name of the student", "course", "instructor", 
        "assessment", "date", "experience", "feelings", "learning", 
        "application", "conclusion"
    ]

    for table in doc.tables:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell_text_lower = cell.text.lower().strip()
                
                for key in priority_keys:
                    val = mappings[key]
                    # Strict mapping rule: 
                    # 1. If it's the topic field, only the topic key can match.
                    # 2. Otherwise, the key must be in the cell text.
                    is_match = False
                    if "journal entry topic" in cell_text_lower:
                        if key == "journal entry topic":
                            is_match = True
                        else:
                            continue # Prevent "learning" from matching "Machine Learning" in topic
                    elif key in cell_text_lower:
                        is_match = True
                        
                    if is_match:
                        if ":" in cell.text:
                            # Safely update label:value in same cell
                            for para in cell.paragraphs:
                                if ":" in para.text:
                                    parts = para.text.split(":", 1)
                                    label_part = parts[0].strip()
                                    
                                    # Capture color before update
                                    original_color = None
                                    if para.runs:
                                        original_color = para.runs[0].font.color.rgb
                                    
                                    # Surgically update runs without clearing paragraph
                                    para.runs[0].text = f"{label_part} :"
                                    apply_run_style(para.runs[0], font_size=12, bold=True)
                                    if original_color:
                                        para.runs[0].font.color.rgb = original_color
                                    
                                    if len(para.runs) > 1:
                                        para.runs[1].text = f" {val}"
                                        apply_run_style(para.runs[1], font_size=12, bold=False)
                                        # Clear other runs if any
                                        for r_idx in range(2, len(para.runs)):
                                            para.runs[r_idx].text = ""
                                    else:
                                        v_run = para.add_run(f" {val}")
                                        apply_run_style(v_run, font_size=12, bold=False)
                        elif i + 1 < len(row.cells):
                            # Fill the target cell next to the label
                            target_cell = row.cells[i+1]
                            if target_cell.paragraphs:
                                first_p = target_cell.paragraphs[0]
                                original_color = None
                                if first_p.runs:
                                    original_color = first_p.runs[0].font.color.rgb
                                
                                safe_update_runs(first_p, str(val))
                                # Re-apply style to all runs
                                for r in first_p.runs:
                                    apply_run_style(r, font_size=12, bold=False)
                                    if original_color:
                                        r.font.color.rgb = original_color
                                
                                # Clear other paragraphs in the cell
                                for p_idx in range(1, len(target_cell.paragraphs)):
                                    for r in target_cell.paragraphs[p_idx].runs:
                                        r.text = ""
                        break

    # 2. Global Pass - Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # Global Spacing Rule: Apply 1.5 spacing inside tables
                    paragraph.paragraph_format.line_spacing = 1.5
                    
                    if not paragraph.text.strip():
                        for run in paragraph.runs:
                            apply_run_style(run, font_size=12, bold=False)
                        continue
                    
                    # Use the non-destructive helper
                    process_label_value_paragraph(paragraph)
                    
                    # Ensure alignment and vertical alignment remain untouched
                    # Vertical alignment is cell-level, and paragraph.alignment was preserved in helper

    # 3. Global Pass - Main Body (Non-table paragraphs)
    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            continue
            
        text = paragraph.text
        # Inline Heading Detection (Non-table)
        is_heading = False
        if len(text) < 80:
            # Rules: ends with ':', OR is UPPERCASE, OR is Title Case, OR starts with list marker
            if (text.endswith(":") or 
                text.isupper() or 
                text.istitle() or 
                (len(text) > 2 and text[0].isalpha() and text[1] == ")")):
                is_heading = True
        
        apply_academic_formatting(paragraph, is_heading=is_heading)

    # 4. FINAL VALIDATION (MANDATORY)
    # Re-verify and Hard Lock "Journal Entry Topic"
    expected_topic = data.get("journal_topic", "").strip()
    for table in doc.tables:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if "journal entry topic" in cell.text.lower():
                    # Check if it's label:value
                    if ":" in cell.text:
                        for p in cell.paragraphs:
                            if ":" in p.text:
                                parts = p.text.split(":", 1)
                                label = parts[0].strip()
                                current_val = parts[1].strip()
                                if current_val != expected_topic:
                                    # Force overwrite
                                    safe_update_runs(p, f"{label} : {expected_topic}")
                                    # Re-apply styling
                                    if p.runs:
                                        apply_run_style(p.runs[0], font_size=12, bold=True)
                                        if len(p.runs) > 1:
                                            # If safe_update_runs put everything in runs[0], 
                                            # we might need to split again if we want bold label
                                            # But safe_update_runs just updates runs[0].
                                            # Let's fix process_label_value_paragraph logic here too if needed.
                                            pass
                    elif i + 1 < len(row.cells):
                        target_cell = row.cells[i+1]
                        if target_cell.text.strip() != expected_topic:
                            if target_cell.paragraphs:
                                safe_update_runs(target_cell.paragraphs[0], expected_topic)
                                for r in target_cell.paragraphs[0].runs:
                                    apply_run_style(r, font_size=12, bold=False)

    doc.save(output_path)
    return output_path
