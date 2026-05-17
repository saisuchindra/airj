from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
import sys
import os

def apply_formatting(doc_path, output_path):
    doc = Document(doc_path)
    
    major_headings = ["experience", "feelings", "learning", "application", "conclusion"]

    def format_run(run, size, bold=None):
        run.font.name = 'Bookman Old Style'
        run.font.size = Pt(size)
        if bold is not None:
            run.bold = bold
        # Ensure the font is applied correctly in all views
        r = run._element.get_or_add_rPr()
        rFonts = r.get_or_add_rFonts()
        for attr in ['w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs']:
            rFonts.set(qn(attr), 'Bookman Old Style')

    def process_paragraph(paragraph, is_in_table=False):
        # Global Spacing Rule: 1.5 for ALL paragraphs
        paragraph.paragraph_format.line_spacing = 1.5

        text = paragraph.text.strip()
        if not text:
            return

        # Check if it's a major heading (case insensitive)
        if text.lower() in major_headings:
            for run in paragraph.runs:
                format_run(run, 14, bold=True)
            return

        # Detect label:value pattern
        if ":" in text and len(text.split(":")[0]) < 40:
            parts = text.split(":", 1)
            label_text = parts[0].strip() + " :"
            value_text = " " + parts[1].strip()
            
            # Capture original first run's color
            original_color = None
            if paragraph.runs:
                original_color = paragraph.runs[0].font.color.rgb
            
            paragraph.clear()
            
            l_run = paragraph.add_run(label_text)
            format_run(l_run, 12, bold=True)
            if original_color:
                l_run.font.color.rgb = original_color
            
            if value_text.strip():
                v_run = paragraph.add_run(value_text)
                format_run(v_run, 12, bold=False)
            return

        # Detect inline headings (only outside tables)
        is_inline_heading = False
        if not is_in_table:
            if text.endswith(":") or (len(text) < 60 and ":" in text and len(text.split(":")[0]) < 30):
                 is_inline_heading = True
            elif len(text) < 50 and not text.endswith(".") and not text.endswith("!") and not text.endswith("?"):
                 is_inline_heading = True

        # Preserve alignment if set; apply default only to body
        if paragraph.alignment is None and not is_in_table:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        if is_inline_heading:
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(6)
            for run in paragraph.runs:
                format_run(run, 14, bold=True)
        else:
            # Main body content
            for run in paragraph.runs:
                format_run(run, 12, bold=False)

    # Process all paragraphs in the document
    for para in doc.paragraphs:
        process_paragraph(para)

    # Process all tables in the document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_paragraph(para, is_in_table=True)

    # Global font setting for any newly added text or defaults
    for style in doc.styles:
        if hasattr(style, 'font'):
            style.font.name = 'Bookman Old Style'

    doc.save(output_path)
    print(f"Formatted document saved to: {output_path}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python refine_document.py <input_path> [output_path]")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else input_file.replace(".docx", "_Refined.docx")
    
    if not os.path.exists(input_file):
        print(f"Error: File {input_file} not found.")
        sys.exit(1)
        
    apply_formatting(input_file, output_file)
