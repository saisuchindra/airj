from docx import Document

def verify_formatting(file_path):
    doc = Document(file_path)
    print(f"Verifying {file_path}...")
    
    for t_idx, table in enumerate(doc.tables):
        print(f"\nTable {t_idx}:")
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, paragraph in enumerate(cell.paragraphs):
                    if paragraph.text.strip():
                        print(f"  Table {t_idx}, Row {r_idx}, Cell {c_idx}, P {p_idx}: '{paragraph.text}'")
                        for run in paragraph.runs:
                            font_size = run.font.size.pt if run.font.size else "Default"
                            print(f"    Run: '{run.text}', Bold: {run.bold}, Size: {font_size}")

if __name__ == "__main__":
    verify_formatting("output/Formatting_Test.docx")
