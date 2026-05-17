"""
Tests for structure-aware DOCX filling.
"""

import unittest
from pathlib import Path

from docx import Document

from utils.document_service import fill_reflective_journal


class TestDocumentService(unittest.TestCase):
    def setUp(self):
        self.base_dir = Path(__file__).parent.parent
        self.tmp_dir = self.base_dir / "output"
        self.tmp_dir.mkdir(parents=True, exist_ok=True)
        self.template_path = self.tmp_dir / "reflective_template_test.docx"
        self.output_path = self.tmp_dir / "reflective_output_test.docx"

        doc = Document()
        table = doc.add_table(rows=5, cols=2)
        table.cell(0, 0).text = "Name of the Student"
        table.cell(1, 0).text = "Course"
        table.cell(2, 0).text = "Instructor"
        table.cell(3, 0).text = "Assessment"
        table.cell(4, 0).text = "Date of Submission"

        doc.add_paragraph("Journal Entry Topic")
        doc.add_paragraph("Old topic")
        doc.add_heading("Experience", level=1)
        doc.add_paragraph("Old experience text")
        doc.add_heading("Feelings", level=1)
        doc.add_paragraph("Old feelings text")
        doc.add_heading("Learning", level=1)
        doc.add_paragraph("Old learning text")
        doc.add_heading("Application", level=1)
        doc.add_paragraph("Old application text")
        doc.add_heading("Conclusion", level=1)
        doc.add_paragraph("Old conclusion text")
        doc.save(str(self.template_path))

    def test_fill_reflective_journal_populates_table_topic_and_sections(self):
        data = {
            "student_name": "Alex Doe",
            "course_name": "B.Tech CSE",
            "instructor": "Dr. Smith",
            "assessment": "Reflective Journal 1",
            "date": "2026-04-26",
            "topic": "Cloud Migration Lessons",
            "experience": "I observed multiple code reviews.",
            "feelings": "I felt challenged but motivated.",
            "learning": "I learned practical review techniques.",
            "application": "I will apply this in team projects.",
            "conclusion": "The experience improved my engineering maturity.",
        }

        output = fill_reflective_journal(
            template_path=str(self.template_path),
            output_path=str(self.output_path),
            data=data,
        )
        self.assertTrue(Path(output).exists())

        generated = Document(output)
        para_text = "\n".join(p.text for p in generated.paragraphs)
        table_text = "\n".join(
            cell.text for table in generated.tables for row in table.rows for cell in row.cells
        )
        all_text = f"{para_text}\n{table_text}"

        self.assertIn(data["student_name"], all_text)
        self.assertIn(data["course_name"], all_text)
        self.assertIn(data["instructor"], all_text)
        self.assertIn(data["assessment"], all_text)
        self.assertIn(data["date"], all_text)
        self.assertIn(data["topic"], all_text)
        self.assertIn(data["experience"], all_text)
        self.assertIn(data["feelings"], all_text)
        self.assertIn(data["learning"], all_text)
        self.assertIn(data["application"], all_text)
        self.assertIn(data["conclusion"], all_text)


if __name__ == "__main__":
    unittest.main()
