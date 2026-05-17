"""
Utility Script: Generate a heading-based reflective journal template.
Run this script once to populate /uploads/reflective_journal_template.docx
"""

import os
from pathlib import Path

from docx import Document

def create_sample_docx():
    base_dir = Path(__file__).parent.parent
    uploads_dir = base_dir / "uploads"
    uploads_dir.mkdir(parents=True, exist_ok=True)
    
    file_path = uploads_dir / "reflective_journal_template.docx"
    
    doc = Document()
    doc.add_heading('Reflective Journal Template', 0)

    doc.add_heading('Title', level=1)
    doc.add_paragraph('Add title text here.')

    doc.add_heading('Experience', level=1)
    doc.add_paragraph('Describe the experience.')

    doc.add_heading('Feelings', level=1)
    doc.add_paragraph('Describe your feelings.')

    doc.add_heading('Learning', level=1)
    doc.add_paragraph('Explain what you learned.')

    doc.add_heading('Application', level=1)
    doc.add_paragraph('Explain how you will apply this learning.')

    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph('Write a concise conclusion.')
    
    doc.save(str(file_path))
    print(f"Sample template successfully created at: {file_path}")

if __name__ == "__main__":
    create_sample_docx()
